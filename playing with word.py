from docx import Document
#create new pocket card document
orig = Document('2018 Pocket Card Gym Wheel.docx')
orig.save('New Pocket Card.docx')
poc = Document('New Pocket Card.docx')
#load in the list of tricks
inf = open('Doc.txt', 'r')
lines = inf.readlines()
inf.close()
#clear the data in the cells (not neccesary for a new doc)
for a in range(2,20):
    for b in range(1,5):
        poc.tables[0].cell(a,b).text = ' '
#retrieve desired list of tricks
tricks = input('list of tricks:\n').split(', ')
#for each trick check if it is in the list of tricks
n=1
for i in tricks:
    n += 1
    for k in lines:
        #split the loaded trick data into the
        #different names of the tricks and the
        #values, structure groups, and difficulties
        m = k.split('//')
        if i.title() in m[0]:
            #if the trick is in the list:
            #write the trick into the pocket card
            trick = m[1].split(', ')
            poc.tables[0].cell(n,1).text = trick[0]
            poc.tables[0].cell(n,2).text = trick[1]
            poc.tables[0].cell(n,3).text = trick[2]
            poc.tables[0].cell(n,4).text = trick[3]
            #stop searching and move on to the next trick
            break
poc.save('New Pocket Card.docx')
